# -*- coding: UTF-8 -*- 
#! /usr/local/bin/python

import sys
sys.path.append('../configure/')
sys.path.append('../model/')
sys.path.append('../dao/')
reload(sys)
sys.setdefaultencoding('utf-8')
import configure
from studentInfoDAO import studentInfoDAO
from answerDAO import answerDAO
from selection_questionDAO import selection_questionDAO
from dict_op import sortDict

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
import datetime

my_font=Pt(10)


def write_doc(setid,reading_answers,listening_answers,reading_score,listening_score,report_path):
    """
    生成报告
    """
    #打开文档
    document = Document()

    #加入不同等级的标题
    document.add_heading(u'托福模考报告单',0)

    #添加总评报告
    document.add_heading(u'总评',1)
    #添加总评表格
    table = document.add_table(rows=2,cols=7,style = u'Table Grid')
    table.autofit = False
    # table.font.size=Pt(10)
    hdr_cells=table.columns[0].cells
    #测试题目
    hdr_cells[0].text=u'Test Number'
    hdr_cells[1].text=str(setid)
    hdr_cells[0].width=Inches(1.2)
    #测试日期
    now_time=datetime.datetime.now().strftime('%Y-%m-%d')
    hdr_cells=table.columns[1].cells
    hdr_cells[0].text=u'Your Test Date'
    hdr_cells[1].text=now_time
    hdr_cells[0].width=Inches(1.2)
    #Reading得分
    hdr_cells=table.columns[2].cells
    hdr_cells[0].text=u'Reading'
    hdr_cells[1].text=str(reading_score)
    #Listening得分
    hdr_cells=table.columns[3].cells
    hdr_cells[0].text=u'Listening'
    hdr_cells[1].text=str(listening_score)
    #Speaking得分
    hdr_cells=table.columns[4].cells
    hdr_cells[0].text=u'Speaking'
    #Writing得分
    hdr_cells=table.columns[5].cells
    hdr_cells[0].text=u'Speaking'
    #总得分
    hdr_cells=table.columns[6].cells
    hdr_cells[0].text=u'Total'

    #添加阅读分数
    document.add_heading(u'阅读',1)
    #创建阅读表格
    table = document.add_table(rows=7,cols=16,style = u'Table Grid')
    table.autofit = False
    table.columns[0].cells[0].width=Inches(0.8)
    table.columns[1].cells[0].width=Inches(1.1)
    table.columns[15].cells[0].width=Inches(0.42)
    for index in range(2,15):
        table.columns[index].cells[0].width=Inches(0.1)
    p=table.cell(1,0).add_paragraph()
    p.add_run(u'Passage 1').font.size=my_font
    p=table.cell(3,0).add_paragraph()
    p.add_run(u'Passage 2').font.size=my_font
    p=table.cell(5,0).add_paragraph()
    p.add_run(u'Passage 3').font.size=my_font
    p=table.cell(1,1).add_paragraph()
    p.add_run(u'Your Answer').font.size=my_font
    p=table.cell(3,1).add_paragraph()
    p.add_run(u'Your Answer').font.size=my_font
    p=table.cell(5,1).add_paragraph()
    p.add_run(u'Your Answer').font.size=my_font
    p=table.cell(2,1).add_paragraph()
    p.add_run(u'Correct Answer').font.size=my_font
    p=table.cell(4,1).add_paragraph()
    p.add_run(u'Correct Answer').font.size=my_font
    p=table.cell(6,1).add_paragraph()
    p.add_run(u'Correct Answer').font.size=my_font
    table.cell(1,0).merge(table.cell(2,0))
    table.cell(3,0).merge(table.cell(4,0))
    table.cell(5,0).merge(table.cell(6,0))
    #制作表头
    for index in range(1,15):
        p=table.cell(0,index+1).add_paragraph()
        p.add_run(str(index)).font.size=my_font
        # table.cell(0,index+1).text=str(index)
    #添加答案
    count=0
    index=0
    for single_reading_answers in reading_answers:
        #每一套阅读题答案
        for single_answer in single_reading_answers:
            student_answer=single_answer.keys()[0]
            official_answer=single_answer.get(student_answer)
            #写入学生答案
            p=table.cell(count*2+1,index+2).add_paragraph()
            p.add_run(student_answer).font.size=my_font
            #写入正确答案
            p=table.cell(count*2+2,index+2).add_paragraph()
            p.add_run(official_answer).font.size=my_font
            index+=1

        index=0
        count+=1


    #添加听力分数
    document.add_heading(u'听力',1)
    #创建阅读表格
    table = document.add_table(rows=13,cols=8,style = u'Table Grid')
    table.autofit = False
    table.columns[0].cells[0].width=Inches(1.2)
    table.columns[1].cells[0].width=Inches(1.2)
    table.cell(1,0).text=u'Conversation 1'
    table.cell(3,0).text=u'Lecture 1'
    table.cell(5,0).text=u'Lecture 2'
    table.cell(7,0).text=u'Conversation 2'
    table.cell(9,0).text=u'Lecture 3'
    table.cell(11,0).text=u'Lecture 4'
    table.cell(1,1).text=u'Your Answer'
    table.cell(3,1).text=u'Your Answer'
    table.cell(5,1).text=u'Your Answer'
    table.cell(7,1).text=u'Your Answer'
    table.cell(9,1).text=u'Your Answer'
    table.cell(11,1).text=u'Your Answer'
    table.cell(2,1).text=u'Correct Answer'
    table.cell(4,1).text=u'Correct Answer'
    table.cell(6,1).text=u'Correct Answer'
    table.cell(8,1).text=u'Correct Answer'
    table.cell(10,1).text=u'Correct Answer'
    table.cell(12,1).text=u'Correct Answer'
    table.cell(1,0).merge(table.cell(2,0))
    table.cell(3,0).merge(table.cell(4,0))
    table.cell(5,0).merge(table.cell(6,0))
    table.cell(7,0).merge(table.cell(8,0))
    table.cell(9,0).merge(table.cell(10,0))
    table.cell(11,0).merge(table.cell(12,0))
    #制作表头
    for index in range(1,7):
        table.cell(0,index+1).text=str(index)
    #添加答案
    count=0
    index=0
    for single_listening_answers in listening_answers:
        #每一套阅读题答案
        for single_answer in single_listening_answers:
            student_answer=single_answer.keys()[0]
            official_answer=single_answer.get(student_answer)
            #写入学生答案
            table.cell(count*2+1,index+2).text=student_answer
            #写入正确答案
            table.cell(count*2+2,index+2).text=official_answer
            index+=1

        index=0
        count+=1

    #添加口语分数
    document.add_heading(u'口语',1)
    #创建口语表格
    table = document.add_table(rows=7,cols=3,style = u'Table Grid')
    table.autofit = False
    table.columns[0].cells[0].width=Inches(0.5)
    table.columns[1].cells[0].width=Inches(0.5)
    table.columns[2].cells[0].width=Inches(3)
    table.cell(1,0).text=u'Task 1'
    table.cell(2,0).text=u'Task 2'
    table.cell(3,0).text=u'Task 3'
    table.cell(4,0).text=u'Task 4'
    table.cell(5,0).text=u'Task 5'
    table.cell(6,0).text=u'Task 6'    
    table.cell(0,1).text=u'Score(1-5)' 
    table.cell(0,2).text=u'Comments'

    #添加写作分数
    document.add_heading(u'写作',1)
    #创建写作表格
    table = document.add_table(rows=3,cols=3,style = u'Table Grid')
    table.autofit = False
    table.columns[0].cells[0].width=Inches(0.5)
    table.columns[1].cells[0].width=Inches(0.5)
    table.columns[2].cells[0].width=Inches(3)
    table.cell(1,0).text=u'Integrated Writing'
    table.cell(2,0).text=u'Independent Writing'  
    table.cell(0,1).text=u'Score(1-5)' 
    table.cell(0,2).text=u'Comments'

    #添加建议
    document.add_heading(u'Suggestion',1)

    #增加分页
    document.add_page_break()

    #保存文件
    document.save(report_path)

if __name__ == '__main__':
    path='/Users/apple/Code/TOP/util/demo.docx'
    official_answer={"L18" : "C", "L19" : "BD", "L14" : "B", "L15" : "C", "L16" : "D", "L17" : "D", "L10" : "ACE", "L11" : "B", "L12" : "C", "L13" : "A", "R42" : "BDE", "R16" : "C", "R17" : "D", "R14" : "BEF", "R15" : "C", "R12" : "B", "R13" : "C", "R10" : "A", "R11" : "A", "R40" : "B", "R18" : "C", "R19" : "B", "L21" : "B", "L20" : "B", "L23" : "C", "L22" : "C", "L25" : "AB", "L24" : "D", "L27" : "C", "L26" : "C", "R34" : "C", "R35" : "D", "R36" : "A", "R37" : "B", "R30" : "D", "R31" : "C", "R32" : "C", "R33" : "C", "R4" : "A", "R5" : "C", "R6" : "D", "R7" : "C", "R39" : "B", "R1" : "B", "R2" : "B", "R3" : "B", "R28" : "ADE", "R8" : "B", "R9" : "B", "L6" : "B", "L29" : "BDE", "L4" : "B", "L5" : "C", "L2" : "C", "L3" : "A", "L1" : "D", "L8" : "D", "L9" : "C", "R24" : "D", "L30" : "B", "L7" : "A", "L32" : "C", "userid" : -1, "L28" : "B", "R38" : "A", "L34" : "D", "R41" : "D", "L33" : "D", "R29" : "A", "L31" : "A", "R27" : "A", "R26" : "A", "R25" : "B", "R23" : "D", "R22" : "A", "R21" : "C", "R20" : "D" }
    reading_test=[{'R1':'C','R12':"A",'R2':'D','R14':'ABC'},{'R15':'B','R27':"C",'R16':'D','R28':'ABC'},{'R29':'B','R41':"C",'R30':'D','R42':'ABC'}]
    listening_test=reading_test
    listening_test.extend(reading_test)
    write_doc(20170603,1,2,20,24,path)